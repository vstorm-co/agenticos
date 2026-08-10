import asyncio
import hashlib
import logging
import re
import shutil
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

import pymupdf
from docx import Document as DOCXDocument
from langchain_text_splitters import MarkdownHeaderTextSplitter, RecursiveCharacterTextSplitter

from app.services.rag.config import (
    LITEPARSE_OFFICE_FORMATS,
    LITEPARSE_PDF_FORMATS,
    LLAMAPARSE_PDF_FORMATS,
    PYMUPDF_PDF_FORMATS,
    DocumentExtensions,
    PdfParser,
    RAGSettings,
)
from app.services.rag.image_describer import BaseImageDescriber
from app.services.rag.models import (
    Document,
    DocumentImage,
    DocumentMetadata,
    DocumentPage,
    DocumentPageChunk,
)

logger = logging.getLogger(__name__)

# Markdown reconstruction wraps a page with no text in an empty fenced block, so
# ``"```text\n\n```"`` is what an unreadable scan looks like once it reaches us.
# It is not whitespace, so a plain `.strip()` keeps it, and the page goes on to
# be embedded as a chunk that says nothing and matches nothing.
_MARKDOWN_NOISE = re.compile(r"```[a-z]*|`|\s+")


def has_indexable_text(content: str) -> bool:
    """Whether a parsed page carries anything worth embedding.

    Public because it is also the honest answer to "did this document parse to
    anything" when the stored chunks are read back for display.
    """
    return bool(_MARKDOWN_NOISE.sub("", content))


class BaseDocumentParser(ABC):
    allowed = [f"{ext.value}" for ext in DocumentExtensions]

    def is_file_existing(self, filepath: Path) -> bool:
        return Path.exists(filepath)

    def is_extension_allowed(self, filepath: Path) -> bool:
        return filepath.suffix.lower() in self.allowed and self.is_file_existing(filepath)

    def get_document_metadata(self, filepath: Path) -> DocumentMetadata:
        content_hash = hashlib.sha256(filepath.read_bytes()).hexdigest()
        return DocumentMetadata(
            filename=filepath.name,
            filesize=filepath.stat().st_size,
            filetype=filepath.suffix.replace(".", ""),
            source_path=str(filepath.resolve()),
            content_hash=content_hash,
        )

    @abstractmethod
    async def parse(self, filepath: Path) -> Document:
        """Parse a file and read its content into a Document object.
        Args:
            filepath: Path to the file to parse.
        Returns:
            Document object with parsed content and metadata.
        """


class TextDocumentParser(BaseDocumentParser):
    """Parser for text-based documents (TXT, MD).
    Uses Python's built-in file reading capabilities to extract
    text content from plain text and Markdown files.
    """

    # Narrower than the base class, which lists everything the pipeline reads.
    # `is_extension_allowed` lower-cases before comparing, so this is also what
    # makes `README.MD` work: the upload endpoint validates a lower-cased
    # extension, and a parser that then compared `filepath.suffix` verbatim
    # refused in a worker the file the API had just accepted.
    allowed = [".txt", ".md"]

    def _parse_text_file(self, filepath: Path) -> Document:
        """Extract raw text from a TXT or MD file.
        Args:
            filepath: Path to the text file.
        Returns:
            Document object with the file content.
        """
        with open(filepath, encoding="utf-8") as f:
            page = DocumentPage(page_num=1, content=f.read())

        return Document(pages=[page], metadata=self.get_document_metadata(filepath))

    async def parse(self, filepath: Path) -> Document:
        """Parse a text file (TXT or MD).

        Args:
            filepath: Path to the text file.

        Returns:
            Document object with parsed content.

        Raises:
            ValueError: If the file extension is not supported.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by TextDocumentParser")

        return self._parse_text_file(filepath)


class DocxDocumentParser(BaseDocumentParser):
    """Parser for DOCX documents using python-docx.

    Extracts text content from Microsoft Word documents by reading
    all paragraphs and joining them with newline characters.
    """

    allowed = [".docx"]

    def _parse_docx_file(self, filepath: Path) -> Document:
        """Extract raw text from the DOCX file.

        Args:
            filepath: Path to the DOCX file.

        Returns:
            Document object with the file content.
        """
        file: Any = DOCXDocument(str(filepath))
        page = DocumentPage(page_num=1, content="\n".join([p.text for p in file.paragraphs]))
        return Document(pages=[page], metadata=self.get_document_metadata(filepath))

    async def parse(self, filepath: Path) -> Document:
        """Parse a DOCX file.

        Args:
            filepath: Path to the DOCX file.

        Returns:
            Document object with parsed content.

        Raises:
            ValueError: If the file is not a DOCX file.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by DocxDocumentParser")

        return self._parse_docx_file(filepath)


class PyMuPDFParser(BaseDocumentParser):
    """Smart PDF parser using PyMuPDF.

    Features:
    - Text extraction with layout preservation (blocks)
    - Table detection -> markdown tables
    - Header/footer detection and removal
    - OCR fallback for scanned pages (optional, requires tesseract)
    - Image extraction for LLM-based description
    - Document metadata (author, title, TOC)
    """

    MIN_TEXT_LENGTH = 50  # below this -> likely a scan, try OCR

    # PDFs only. The base class allows what the *pipeline* reads, which includes
    # the three formats the Python-native parsers take; this parser is never
    # routed them, and saying so here keeps `is_extension_allowed` honest.
    allowed = sorted(PYMUPDF_PDF_FORMATS)

    def __init__(self, enable_ocr: bool = False, image_describer: Any = None):
        self.enable_ocr = enable_ocr
        self._image_describer = image_describer

    def _detect_repeated_content(self, doc: Any) -> set[str]:
        """Detect headers/footers -- text appearing on >70% of pages."""
        if len(doc) < 3:
            return set()
        text_counts: dict[str, int] = {}
        for page in doc:
            for b in page.get_text("blocks"):
                if b[6] != 0:  # skip image blocks
                    continue
                y_ratio = b[1] / page.rect.height if page.rect.height else 0
                if y_ratio < 0.15 or y_ratio > 0.85:
                    text = b[4].strip()
                    if text and len(text) < 200:
                        text_counts[text] = text_counts.get(text, 0) + 1
        threshold = len(doc) * 0.7
        return {t for t, c in text_counts.items() if c >= threshold}

    def _extract_text(self, page: Any, repeated: set[str]) -> str:
        """Extract text blocks, filtering headers/footers."""
        texts = []
        for b in page.get_text("blocks"):
            if b[6] != 0:  # skip image blocks
                continue
            text = b[4].strip()
            if text and text not in repeated:
                texts.append(text)
        return str("\n\n".join(texts))

    def _extract_tables(self, page: Any) -> str:
        """Extract tables as markdown."""
        try:
            tables = page.find_tables()
            if not tables or not tables.tables:
                return ""
            parts = []
            for table in tables.tables:
                df = table.to_pandas()
                if not df.empty:
                    parts.append(df.to_markdown(index=False))
            return "\n\n".join(parts)
        except Exception:
            return ""

    def _ocr_page(self, page: Any, image_describer: Any = None) -> str:
        """OCR a scanned page by rendering it as image and sending to LLM vision."""
        if not image_describer:
            return ""
        try:
            pix = page.get_pixmap(dpi=200)
            image_bytes = pix.tobytes("png")
            loop = asyncio.new_event_loop()
            try:
                return str(
                    loop.run_until_complete(image_describer.describe(image_bytes, "image/png"))
                )
            finally:
                loop.close()
        except Exception as e:
            logger.warning("LLM OCR failed for page %d: %s", page.number + 1, e)
            return ""

    def _extract_images(self, doc: Any, page: Any) -> list["DocumentImage"]:
        """Extract images from page for LLM description."""
        images = []
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                base = doc.extract_image(xref)
                if base and base["image"] and len(base["image"]) > 1000:
                    ext = base.get("ext", "png")
                    mime_map = {"png": "image/png", "jpeg": "image/jpeg", "jpg": "image/jpeg"}
                    images.append(
                        DocumentImage(
                            page_num=page.number + 1,
                            image_bytes=base["image"],
                            mime_type=mime_map.get(ext, f"image/{ext}"),
                        )
                    )
            except Exception:
                # One unreadable embedded image should not fail a whole
                # document, but silently dropping it is how "the PDF ingested
                # but the diagram is missing" becomes unexplainable.
                logger.warning("Skipping an unreadable image on page %d", page.number + 1)
        return images

    def _parse_pdf_file(self, filepath: Path) -> Document:
        """Parse PDF with smart extraction pipeline."""
        doc: Any = pymupdf.open(filepath)  # type: ignore[no-untyped-call]

        meta = doc.metadata or {}
        toc = doc.get_toc()

        repeated = self._detect_repeated_content(doc)

        pages = []
        for page in doc:
            text = self._extract_text(page, repeated)

            tables_md = self._extract_tables(page)
            if tables_md:
                text = text + "\n\n" + tables_md if text.strip() else tables_md

            if self.enable_ocr and len(text.strip()) < self.MIN_TEXT_LENGTH:
                ocr_text = self._ocr_page(page, self._image_describer)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    logger.info("OCR fallback used for page %d", page.number + 1)
            images = self._extract_images(doc, page)

            pages.append(
                DocumentPage(
                    page_num=page.number + 1,
                    content=text,
                    images=images,
                )
            )

        doc.close()

        additional: dict[str, Any] = {}
        if meta.get("title"):
            additional["pdf_title"] = meta["title"]
        if meta.get("author"):
            additional["pdf_author"] = meta["author"]
        if toc:
            additional["toc"] = [{"level": t[0], "title": t[1], "page": t[2]} for t in toc[:20]]

        doc_meta = self.get_document_metadata(filepath)
        if additional:
            doc_meta.additional_info = {**(doc_meta.additional_info or {}), **additional}

        return Document(pages=pages, metadata=doc_meta)

    async def parse(self, filepath: Path) -> Document:
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by PyMuPDFParser")
        return self._parse_pdf_file(filepath)


class LlamaParseParser(BaseDocumentParser):
    """Advanced document parser using LlamaParse cloud API.

    Provides AI-powered document parsing with support for 130+ formats
    including PDF, DOCX, PPTX, XLSX, images (OCR), and more.
    Returns markdown-formatted content.
    """

    allowed = sorted(LLAMAPARSE_PDF_FORMATS)

    def __init__(self, api_key: str, tier: str = "agentic"):
        """Initialize the LlamaParse parser.

        Args:
            api_key: LlamaCloud API key for authentication.
            tier: Parsing tier (fast, cost_effective, agentic, agentic_plus).
        """
        from llama_cloud import AsyncLlamaCloud

        self.parser = AsyncLlamaCloud(api_key=api_key)
        self.tier = tier

    async def parse(self, filepath: Path) -> Document:
        """Parse a document using LlamaParse.

        Supports PDF, DOCX, PPTX, XLSX, images, and many more formats.
        See https://developers.llamaindex.ai/python/cloud/llamaparse/supported_document_types/

        Args:
            filepath: Path to the file to parse.

        Returns:
            Document object with parsed markdown content.

        Raises:
            ValueError: If the file extension is not supported.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by LlamaParse")

        file_obj = await self.parser.files.create(file=filepath, purpose="parse")
        result = await self.parser.parsing.parse(
            file_id=file_obj.id,
            tier=self.tier,
            version="latest",
            expand=["text", "markdown"],
        )
        pages = []
        for page in result.markdown.pages:
            pages.append(DocumentPage(page_num=page.page_number, content=page.markdown))

        return Document(pages=pages, metadata=self.get_document_metadata(filepath))


class LiteParseParser(BaseDocumentParser):
    """Document parser using LiteParse -- fast, local, layout-aware parsing.

    LiteParse reads PDFs through PDFium and converts everything else to PDF
    first: images natively in its Rust core, office documents by shelling out to
    LibreOffice. Nothing leaves the machine and there is no API key, which is
    what makes it the useful middle between PyMuPDF and LlamaParse.

    Configuration is passed to the **constructor**, not to `parse()`, which
    takes the file and nothing else. The previous version of this class passed
    `ocr_enabled`/`ocr_server_url`/`ocr_language`/`timeout` as keyword
    arguments to `parse()` and read `page.pageNum` (the Node binding's
    spelling; Python exposes `page_num`). Both are hard errors, so every
    LiteParse ingestion raised `TypeError` on the first call - and since
    `TypeError` was outside the `except` clauses below, it surfaced as an
    unhandled worker crash rather than a parse failure. The parser was
    selectable in the UI and had never once run.

    Requires: `pip install liteparse` (ships its own native binary; there is
    no Node.js CLI to pre-install). Office formats additionally need LibreOffice
    on PATH -- see :meth:`libreoffice_available`.
    """

    allowed = sorted(LITEPARSE_PDF_FORMATS)

    # Probed once per process. `shutil.which` is a handful of stat calls, but
    # this is consulted per document to explain a failure, and the answer cannot
    # change without the container being replaced.
    _libreoffice: bool | None = None

    def __init__(
        self,
        *,
        enable_ocr: bool = True,
        auto_ocr: bool = True,
        ocr_server_url: str | None = None,
        ocr_language: str = "eng",
        timeout_seconds: float = 600.0,
        output_format: str = "markdown",
        dpi: float = 150.0,
        max_pages: int = 1000,
        num_workers: int | None = None,
    ) -> None:
        self.enable_ocr = enable_ocr
        self.auto_ocr = auto_ocr
        self.ocr_language = ocr_language
        self.timeout_seconds = timeout_seconds
        self.output_format = output_format
        self._options: dict[str, Any] = {
            "ocr_server_url": ocr_server_url,
            "ocr_language": ocr_language,
            "output_format": output_format,
            "dpi": dpi,
            "max_pages": max_pages,
            "num_workers": num_workers,
            # Progress output goes to stdout, where it interleaves with the
            # worker's logs and belongs to no request in particular.
            "quiet": True,
        }

    @classmethod
    def libreoffice_available(cls) -> bool:
        """Whether office documents can be converted on this machine.

        Mirrors liteparse's own discovery order (`find_libre_office_command` in
        `conversion.rs`): the two command names first, then the macOS bundle.
        """
        if cls._libreoffice is None:
            cls._libreoffice = bool(
                shutil.which("libreoffice")
                or shutil.which("soffice")
                or Path("/Applications/LibreOffice.app/Contents/MacOS/soffice").exists()
            )
        return cls._libreoffice

    def _build(self, *, ocr: bool) -> Any:
        from liteparse import LiteParse

        return LiteParse(ocr_enabled=ocr, **self._options)

    def _needs_ocr(self, filepath: Path) -> bool:
        """Ask the cheap text-layer pass whether OCR is worth running.

        This is the single biggest saving LiteParse offers. OCR dominates the
        cost of a parse, and most PDFs are born digital and need none of it;
        `is_complex` reads only the text layer and reports per page. Running it
        first turns "OCR every page of every document" into "OCR the documents
        that are actually scans".

        A failure here is not a parse failure - if the check cannot answer, the
        honest thing is to parse with OCR on rather than to refuse.
        """
        try:
            stats = self._build(ocr=False).is_complex(str(filepath))
        except Exception:
            logger.warning("LiteParse: complexity check failed for %s", filepath.name)
            return True
        return any(page.needs_ocr for page in stats)

    async def parse(self, filepath: Path) -> Document:
        """Parse a document using LiteParse.

        Catches LiteParse exceptions and re-raises as RuntimeError so callers
        can surface them as a structured ingestion failure instead of an
        opaque subprocess error.
        """

        from liteparse.types import ParseError  # type: ignore[import-not-found]

        if filepath.suffix.lower() in LITEPARSE_OFFICE_FORMATS and not self.libreoffice_available():
            raise RuntimeError(
                f"LiteParse converts {filepath.suffix} through LibreOffice, which is not "
                f"installed. Install it (`apt-get install libreoffice`), or choose a parser "
                f"that reads {filepath.suffix} directly."
            )

        ocr = self.enable_ocr and (not self.auto_ocr or self._needs_ocr(filepath))

        try:
            # The Python binding is synchronous, and a parse is CPU-bound work
            # measured in seconds -- doing it on the event loop stalls every
            # other request on this worker.
            #
            # `wait_for` bounds how long *we* wait, not how long the thread
            # runs: a cancelled `to_thread` leaves its thread working to
            # completion. It is a ceiling on the request, not on the machine,
            # which is why `max_pages` is the setting that actually bounds cost.
            result = await asyncio.wait_for(
                asyncio.to_thread(self._build(ocr=ocr).parse, str(filepath)),
                timeout=self.timeout_seconds,
            )
        except FileNotFoundError as e:
            raise RuntimeError(f"LiteParse: file not found: {filepath}") from e
        except TimeoutError as e:
            raise RuntimeError(
                f"LiteParse: parse timed out after {self.timeout_seconds}s for {filepath.name}"
            ) from e
        except ParseError as e:
            raise RuntimeError(f"LiteParse: parse failed for {filepath.name}: {e}") from e

        # `page.markdown` is populated only in markdown mode; `page.text` always
        # holds the spatial-layout rendering, so it is the fallback rather than
        # a competing choice.
        pages: list[DocumentPage] = []
        for page in result.pages:
            content = page.markdown if self.output_format == "markdown" and page.markdown else ""
            content = content or page.text
            if has_indexable_text(content):
                pages.append(DocumentPage(page_num=page.page_num, content=content))

        return Document(
            pages=pages,
            metadata=self.get_document_metadata(filepath),
        )


class PdfParserFactory:
    """Builds the parser a collection's stored configuration asks for."""

    @staticmethod
    def create(
        parser_name: str, settings: RAGSettings | None = None, image_describer: Any = None
    ) -> BaseDocumentParser:
        if parser_name == "llamaparse":
            if not settings or not settings.pdf_parser.api_key:
                raise ValueError("LlamaParse requires LLAMAPARSE_API_KEY to be set")
            return LlamaParseParser(
                api_key=settings.pdf_parser.api_key,
                tier=settings.pdf_parser.tier,
            )
        if parser_name == "liteparse":
            pdf = settings.pdf_parser if settings else PdfParser()
            return LiteParseParser(
                enable_ocr=settings.enable_ocr if settings else True,
                auto_ocr=pdf.liteparse_auto_ocr,
                ocr_server_url=pdf.liteparse_ocr_server_url,
                ocr_language=pdf.liteparse_ocr_language,
                timeout_seconds=pdf.liteparse_timeout_seconds,
                output_format=pdf.liteparse_output_format,
                dpi=pdf.liteparse_dpi,
                max_pages=pdf.liteparse_max_pages,
            )
        return PyMuPDFParser(
            enable_ocr=settings.enable_ocr if settings else False,
            image_describer=image_describer,
        )


class DocumentProcessor:
    """Orchestrates parsing and chunking of files into Document objects.

    Manages the document processing pipeline:
    1. Route to appropriate parser based on file extension
    2. Parse document content
    3. Chunk document pages using RecursiveCharacterTextSplitter
    Supported file types:
    - TXT, MD: TextDocumentParser (Python native)
    - DOCX: DocxDocumentParser (Python native)
    - PDF: PdfParserFactory selects PyMuPDF, LlamaParse, or LiteParse at runtime
    """

    def __init__(self, settings: RAGSettings, image_describer: BaseImageDescriber | None = None):
        """Initialize the document processor.

        Args:
            settings: RAG configuration settings.
            image_describer: The model that reads pictures found inside a
                document, or `None` to index them as nothing. Injected rather
                than built here: which model that is belongs to the collection
                the document is being ingested into, and its credential is
                sealed for that collection's organization - neither of which
                this class can see. See
                :class:`app.services.ingestion_config.IngestionConfigService`.
        """
        self.settings = settings
        self.splitter = self._create_splitter(settings)

        self.text_parser = TextDocumentParser()
        self.docx_parser = DocxDocumentParser()
        self.image_describer = image_describer
        self.pdf_parser = PdfParserFactory.create(
            parser_name=settings.pdf_parser.method,
            settings=settings,
            image_describer=self.image_describer,
        )

    @staticmethod
    def _create_splitter(settings: RAGSettings) -> Any:
        """Create text splitter based on chunking strategy."""
        strategy = settings.chunking_strategy

        if strategy == "markdown":
            return MarkdownHeaderTextSplitter(
                headers_to_split_on=[
                    ("#", "h1"),
                    ("##", "h2"),
                    ("###", "h3"),
                ],
                strip_headers=False,
            )

        if strategy == "fixed":
            return RecursiveCharacterTextSplitter(
                chunk_size=settings.chunk_size,
                chunk_overlap=settings.chunk_overlap,
                length_function=len,
                separators=["\n"],
            )

        return RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )

    async def _describe_images(self, document: Document) -> None:
        """Generate text descriptions for all images in document pages."""
        if self.image_describer is None:
            return
        for page in document.pages:
            if not page.images:
                continue
            for image in page.images:
                image.description = await self.image_describer.describe(
                    image.image_bytes, image.mime_type
                )
            img_descriptions = [
                f"[Image: {img.description}]" for img in page.images if img.description
            ]
            if img_descriptions:
                page.content = f"{page.content}\n\n{chr(10).join(img_descriptions)}"

    async def process_file(self, filepath: Path) -> Document:
        """Main entry point: filepath -> Document with chunks.

        Args:
            filepath: Path to the file to process.

        Returns:
            Document object with parsed pages and chunked content.

        Raises:
            ValueError: If the configured parser cannot read the file type.
        """
        suffix = filepath.suffix.lower()
        if suffix in (".txt", ".md"):
            document = await self.text_parser.parse(filepath)
        elif suffix == ".docx":
            document = await self.docx_parser.parse(filepath)
        elif suffix in self.pdf_parser.allowed:
            # Not only `.pdf`: LlamaParse is routed the spreadsheets, decks and
            # images it accepts. Asking the parser what it takes, rather than
            # naming extensions here, is what keeps this branch in step with
            # `PARSER_FORMATS` - the set the upload endpoint validated against
            # before the file ever reached a worker.
            document = await self.pdf_parser.parse(filepath)
        else:
            raise ValueError(
                f"Unsupported file type: {suffix}. "
                f"{type(self.pdf_parser).__name__} reads {', '.join(self.pdf_parser.allowed)}"
            )
        await self._describe_images(document)

        pages = document.pages

        chunked_pages: list[DocumentPageChunk] = []
        is_markdown_splitter = self.settings.chunking_strategy == "markdown"
        for page in pages:
            if is_markdown_splitter:
                # MarkdownHeaderTextSplitter returns Document objects
                md_docs = self.splitter.split_text(page.content)
                chunks = [doc.page_content for doc in md_docs]
            else:
                chunks = self.splitter.split_text(page.content)
            for chunk_num, chunk in enumerate(chunks):
                chunked_pages.append(
                    DocumentPageChunk(
                        chunk_content=chunk,
                        chunk_num=chunk_num,
                        parent_doc_id=document.id,
                        **page.model_dump(exclude={"parent_doc_id"}),
                    )
                )

        if not chunked_pages:
            # A document that yields nothing is indistinguishable, once stored,
            # from one that ingested fine and simply never matches a search. The
            # overwhelmingly common cause is a scan parsed with OCR off, so the
            # message names that rather than making somebody guess.
            raise ValueError(
                f"{filepath.name} produced no indexable text. If it is a scan or an image, "
                f"turn OCR on for this collection; {type(self.pdf_parser).__name__} read it "
                f"with OCR {'on' if self.settings.enable_ocr else 'off'}."
            )

        document.chunked_pages = chunked_pages
        return document
